import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import os
import sys
import re
import tempfile
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt

# -------------------------------------------------------------------
# Your company details (customise once)
# -------------------------------------------------------------------
COMPANY_NAME = "Your Company Name"
COMPANY_ADDRESS = "123 Chemical Lane, Industrial City"
COMPANY_PHONE = "+1 555 123 4567"
COMPANY_WEBSITE = "www.yourcompany.com"

# -------------------------------------------------------------------
# Helper for PyInstaller bundled resources
# -------------------------------------------------------------------
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# -------------------------------------------------------------------
# PDF to DOCX conversion (keeps formatting)
# -------------------------------------------------------------------
def pdf_to_docx(pdf_path, docx_path):
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    return docx_path

# -------------------------------------------------------------------
# Remove old supplier sections from the DOCX body
# -------------------------------------------------------------------
def delete_supplier_section(doc):
    """Delete all paragraphs and tables from the start of the document
    up to the first safety heading (after Section 1)."""
    
    # Headings that mark the end of Section 1 (supplier identification)
    stop_patterns = [
        r'\bSECTION\s*2\b',
        r'\bHAZARDS?\s*IDENTIFICATION\b',
        r'\bCOMPOSITION\s*\/?\s*INFORMATION\s+ON\s+INGREDIENTS\b',
        r'2\.\s+HAZARDS?\s*IDENTIFICATION',
        r'2\.\s+COMPOSITION'
    ]
    
    # Gather all block-level elements in order
    body = doc.element.body
    elements_to_delete = []
    stop = False
    
    for child in body:
        if stop:
            break
        # Check if child is a paragraph or table
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl'):
            # Get text of this element (for paragraphs)
            text = ''
            if tag == 'p':
                # Extract text from paragraph
                text = child.text or ''
                for r in child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if r.text:
                        text += r.text
            else:
                # For tables, we can't easily get text in one line; we'll use the whole table's text later
                # Better to check the table's first row or just include tables in stop detection via paragraphs inside
                # We'll check the first paragraph inside the table for simplicity
                para = child.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
                if para is not None:
                    text = para.text or ''
                    for r in para.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if r.text:
                            text += r.text
            
            # Check if text matches a stop pattern
            for pat in stop_patterns:
                if re.search(pat, text, re.IGNORECASE):
                    stop = True
                    break
            
            if not stop:
                elements_to_delete.append(child)
    
    # Remove the identified elements
    for elem in elements_to_delete:
        body.remove(elem)

# -------------------------------------------------------------------
# Clean old headers/footers that appear as repeating text in body
# -------------------------------------------------------------------
def strip_repeating_header_lines(doc, supplier_names):
    """Remove paragraphs that look like per-page headers/footers containing
    supplier details, page numbers, or web addresses."""
    removal_pattern = re.compile(
        r'(Page\s+\d+\s+of\s+\d+|' +
        '|'.join(re.escape(name) for name in supplier_names) +
        r'|www\.|http|Tel|Fax|Phone|Emergency\s*Tel)',
        re.IGNORECASE
    )
    for para in doc.paragraphs:
        if removal_pattern.search(para.text):
            # Remove the paragraph element completely
            p = para._element
            p.getparent().remove(p)
    # Also check tables (rarely used for headers)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if removal_pattern.search(para.text):
                        p = para._element
                        p.getparent().remove(p)

# -------------------------------------------------------------------
# Replace product name in the remaining body
# -------------------------------------------------------------------
def replace_product_name(doc, new_name):
    """Find 'Product name' or 'Product identifier' and change the next run
    or the same paragraph to the new name."""
    patterns = [
        r'(Product\s*name\s*:?\s*)',
        r'(Product\s*identifier\s*:?\s*)',
        r'(1\.1\s*Product\s*identifier\s*)'
    ]
    for para in doc.paragraphs:
        for pat in patterns:
            match = re.search(pat, para.text, re.IGNORECASE)
            if match:
                # We'll replace the whole paragraph text with "Product name: new_name"
                para.text = match.group(1) + new_name
                return

# -------------------------------------------------------------------
# Apply your company header/footer to all sections
# -------------------------------------------------------------------
def apply_company_header_footer(doc):
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].clear()
        run = header.paragraphs[0].add_run(f"{COMPANY_NAME} – {COMPANY_ADDRESS} – {COMPANY_PHONE}")
        run.font.size = Pt(9)
        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0].clear()
        run = footer.paragraphs[0].add_run(f"Confidential – {COMPANY_WEBSITE}")
        run.font.size = Pt(8)

# -------------------------------------------------------------------
# Main processing pipeline
# -------------------------------------------------------------------
def process_sds(pdf_path, output_path, product_name=None, progress_callback=None):
    if not product_name:
        product_name = os.path.splitext(os.path.basename(pdf_path))[0]

    # Temp file for intermediate DOCX
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        temp_docx = tmp.name

    try:
        if progress_callback: progress_callback("Converting PDF to DOCX (preserving formatting)...")
        pdf_to_docx(pdf_path, temp_docx)

        if progress_callback: progress_callback("Removing supplier identification...")
        doc = Document(temp_docx)
        
        # Delete Section 1 (supplier block)
        delete_supplier_section(doc)
        
        # Remove repeating header/footer lines (common supplier names)
        supplier_names = [
            "Sigma-Aldrich", "Fisher Scientific", "Merck", "BASF", "Dow", "DuPont",
            "Thermo Fisher", "VWR", "Avantor", "Acros Organics"
        ]  # extend as needed
        strip_repeating_header_lines(doc, supplier_names)
        
        # Replace product name
        replace_product_name(doc, product_name)
        
        # Apply your company header/footer
        apply_company_header_footer(doc)

        if progress_callback: progress_callback("Saving final document...")
        doc.save(output_path)

    finally:
        # Clean up temp file
        if os.path.exists(temp_docx):
            os.unlink(temp_docx)

    return product_name

# -------------------------------------------------------------------
# GUI (same as before, but with preview showing DOCX conversion note)
# -------------------------------------------------------------------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("SDS Rebrander (Format-Preserving)")
        self.geometry("650x550")
        self.resizable(True, True)
        self.build_main_ui()

    def build_main_ui(self):
        tk.Label(self, text="SDS Rebrander", font=("Arial", 16, "bold")).pack(pady=10)
        tk.Label(self, text="Supplier details are removed, formatting is kept.\n"
                           "Product name is replaced, company header/footer added.",
                 wraplength=500, justify="left").pack(pady=5)

        frame = tk.Frame(self)
        frame.pack(pady=5)
        self.file_label = tk.Label(frame, text="No PDF selected", fg="gray", width=45, anchor="w")
        self.file_label.pack(side="left", padx=5)
        tk.Button(frame, text="Choose PDF...", command=self.select_pdf).pack(side="left")

        name_frame = tk.Frame(self)
        name_frame.pack(pady=5)
        tk.Label(name_frame, text="Product name (blank = use filename):").pack(side="left")
        self.product_name_var = tk.StringVar()
        tk.Entry(name_frame, textvariable=self.product_name_var, width=30).pack(side="left", padx=5)

        self.convert_btn = tk.Button(self, text="Convert", state="disabled", command=self.convert)
        self.convert_btn.pack(pady=10)

        self.status = tk.Label(self, text="", fg="blue")
        self.status.pack(pady=5)

        # Preview area (shows first bit of converted text for diagnostics)
        tk.Label(self, text="Extracted text preview (first 2000 chars):", anchor="w").pack(pady=(10,0))
        self.preview = scrolledtext.ScrolledText(self, height=10, width=70, state="disabled")
        self.preview.pack(padx=10, pady=5)

        self.pdf_path = None

    def select_pdf(self):
        path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if path:
            self.pdf_path = path
            self.file_label.config(text=os.path.basename(path))
            self.convert_btn.config(state="normal")
            base = os.path.splitext(os.path.basename(path))[0]
            self.product_name_var.set(base)
            # Quick preview using old PyPDF2 text extraction (just for diagnostics)
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(path)
                text = ""
                for page in reader.pages[:2]:  # first two pages
                    t = page.extract_text()
                    if t: text += t + "\n"
                self.preview.config(state="normal")
                self.preview.delete(1.0, tk.END)
                if text.strip():
                    self.preview.insert(tk.END, text[:2000])
                else:
                    self.preview.insert(tk.END, "⚠️ Scanned PDF detected. Conversion will still work (using layout), but text may be limited.")
                self.preview.config(state="disabled")
            except Exception:
                pass

    def convert(self):
        if not self.pdf_path:
            return
        output_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                   filetypes=[("Word Document", "*.docx")],
                                                   initialfile=self.product_name_var.get() + ".docx")
        if not output_path:
            return
        self.convert_btn.config(state="disabled")
        self.status.config(text="Working...")
        self.update()

        try:
            product_name = self.product_name_var.get().strip() or None
            process_sds(self.pdf_path, output_path, product_name,
                        progress_callback=lambda msg: self.status.config(text=msg))
            self.status.config(text=f"Done! Saved to {output_path}")
            messagebox.showinfo("Success", f"Formatted file saved to:\n{output_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Conversion failed:\n{str(e)}")
            self.status.config(text="Error occurred.")
        finally:
            self.convert_btn.config(state="normal")

if __name__ == "__main__":
    app = App()
    app.mainloop()
