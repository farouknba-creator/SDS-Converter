import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import json
import os
import sys
import requests
from docx import Document
import pdfplumber   # better table/text extraction than PyPDF2

# -------------------------------------------------------------------
# CONFIGURATION
# -------------------------------------------------------------------
DEEPSEEK_API_KEY = ""   # paste your key, or leave blank to be prompted
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
MODEL = "deepseek-chat"

# -------------------------------------------------------------------
# The extraction prompt – edit this JSON structure to match your template
# -------------------------------------------------------------------
EXTRACTION_PROMPT = """
You are a chemical safety expert. Extract the following information from the SDS text below.
Return ONLY a valid JSON object with these exact keys (use empty string or empty list if not found):

{
  "product_name": "...",
  "supplier_name": "...",
  "supplier_code": "...",
  "cas_numbers": ["...", "..."],
  "hazard_statements": ["H...", "..."],
  "composition": "...",
  "first_aid_measures": "...",
  "fire_fighting": "...",
  "accidental_release": "...",
  "handling_storage": "...",
  "exposure_controls": "...",
  "physical_chemical": "...",
  "stability_reactivity": "...",
  "toxicological": "...",
  "ecological": "...",
  "disposal": "...",
  "transport": "...",
  "regulatory": "...",
  "other": "..."
}

SDS text:
---
{sds_text}
---
JSON:
"""

# -------------------------------------------------------------------
# Helpers
# -------------------------------------------------------------------
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def load_mapping():
    map_path = resource_path("mapping.json")
    if not os.path.exists(map_path):
        return {}
    with open(map_path, "r", encoding="utf-8") as f:
        return json.load(f)

def extract_text_from_pdf(pdf_path):
    """Use pdfplumber for better extraction (handles tables well)."""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

def call_deepseek(prompt, api_key):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": "You are a precise JSON-only extractor. Never include explanations."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.0,
        "max_tokens": 3000
    }
    resp = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]

def extract_structured_data(sds_text, api_key):
    prompt = EXTRACTION_PROMPT.format(sds_text=sds_text)
    raw = call_deepseek(prompt, api_key)
    # Clean possible markdown fences
    raw = raw.strip().removeprefix("```json").removeprefix("```").rstrip("```").strip()
    return json.loads(raw)

def map_internal_code(data, mapping):
    supplier = data.get("supplier_name", "")
    code = data.get("supplier_code", "")
    internal = mapping.get(supplier) or mapping.get(code) or code
    data["internal_code"] = internal
    return data

def fill_template(data, output_path):
    template_path = resource_path("template.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError("template.docx not found")

    doc = Document(template_path)

    # Replace in paragraphs
    for para in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in para.text:
                if isinstance(value, list):
                    value = ", ".join(value)
                if para.text.strip() == placeholder:
                    # If the placeholder is the whole paragraph, keep the style
                    para.clear()
                    para.add_run(str(value))
                else:
                    para.text = para.text.replace(placeholder, str(value))

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in para.text:
                            if isinstance(value, list):
                                value = ", ".join(value)
                            if para.text.strip() == placeholder:
                                para.clear()
                                para.add_run(str(value))
                            else:
                                para.text = para.text.replace(placeholder, str(value))

    doc.save(output_path)

# -------------------------------------------------------------------
# GUI
# -------------------------------------------------------------------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("SDS Converter – Field Mapping")
        self.geometry("600x500")
        self.resizable(True, True)
        self.api_key = DEEPSEEK_API_KEY

        if not self.api_key:
            self.show_api_key_prompt()
        else:
            self.build_main_ui()

    def show_api_key_prompt(self):
        for widget in self.winfo_children():
            widget.destroy()
        tk.Label(self, text="Enter your DeepSeek API key:", font=("Arial", 12)).pack(pady=20)
        self.api_entry = tk.Entry(self, width=50, show="*")
        self.api_entry.pack(pady=5)
        tk.Button(self, text="Save & Continue", command=self.save_api_key).pack(pady=10)

    def save_api_key(self):
        key = self.api_entry.get().strip()
        if not key:
            messagebox.showerror("Error", "API key cannot be empty.")
            return
        self.api_key = key
        for widget in self.winfo_children():
            widget.destroy()
        self.build_main_ui()

    def build_main_ui(self):
        tk.Label(self, text="SDS Converter", font=("Arial", 16, "bold")).pack(pady=10)
        tk.Label(self, text="Extracts standard fields and fills your company template.\n"
                           "All formatting is from the template, not the source.",
                 wraplength=500, justify="left").pack(pady=5)

        # File selection
        frame = tk.Frame(self)
        frame.pack(pady=5)
        self.file_label = tk.Label(frame, text="No PDF selected", fg="gray", width=45, anchor="w")
        self.file_label.pack(side="left", padx=5)
        tk.Button(frame, text="Choose PDF...", command=self.select_pdf).pack(side="left")

        # Product name override (optional)
        name_frame = tk.Frame(self)
        name_frame.pack(pady=5)
        tk.Label(name_frame, text="Product name (blank = extract from PDF):").pack(side="left")
        self.product_name_var = tk.StringVar()
        tk.Entry(name_frame, textvariable=self.product_name_var, width=30).pack(side="left", padx=5)

        # Convert button
        self.convert_btn = tk.Button(self, text="Convert", state="disabled", command=self.convert)
        self.convert_btn.pack(pady=10)

        self.status = tk.Label(self, text="", fg="blue")
        self.status.pack(pady=5)

        # Preview area
        tk.Label(self, text="Extracted text preview (first 2000 chars):", anchor="w").pack(pady=(10,0))
        self.preview = scrolledtext.ScrolledText(self, height=10, width=70, state="disabled")
        self.preview.pack(padx=10, pady=5)

        self.pdf_path = None

    def select_pdf(self):
        path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if path:
            self.pdf_path = path
            self.file_label.config(text=os.path.basename(path))
            self.convert_btn.config(state="normal")
            # Preview
            try:
                raw = extract_text_from_pdf(path)
                self.preview.config(state="normal")
                self.preview.delete(1.0, tk.END)
                if raw:
                    self.preview.insert(tk.END, raw[:2000])
                else:
                    self.preview.insert(tk.END, "⚠️ No text found – PDF may be scanned.")
                self.preview.config(state="disabled")
            except Exception as e:
                self.preview.config(state="normal")
                self.preview.delete(1.0, tk.END)
                self.preview.insert(tk.END, f"Error: {e}")
                self.preview.config(state="disabled")

    def convert(self):
        if not self.pdf_path:
            return
        output_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                   filetypes=[("Word Document", "*.docx")],
                                                   initialfile=os.path.splitext(os.path.basename(self.pdf_path))[0] + "_converted.docx")
        if not output_path:
            return

        self.convert_btn.config(state="disabled")
        self.status.config(text="Extracting text...")
        self.update()

        try:
            raw_text = extract_text_from_pdf(self.pdf_path)
            if not raw_text:
                raise ValueError("No text extracted – scanned PDF? Try OCR first.")

            self.status.config(text="Calling DeepSeek to extract fields...")
            self.update()
            data = extract_structured_data(raw_text, self.api_key)

            # Override product name if provided
            manual_name = self.product_name_var.get().strip()
            if manual_name:
                data["product_name"] = manual_name

            # Map internal code
            mapping = load_mapping()
            data = map_internal_code(data, mapping)

            self.status.config(text="Filling template...")
            self.update()
            fill_template(data, output_path)

            self.status.config(text=f"Done! Saved to {output_path}")
            messagebox.showinfo("Success", f"Document saved to:\n{output_path}")

        except Exception as e:
            messagebox.showerror("Error", str(e))
            self.status.config(text="Error occurred.")
        finally:
            self.convert_btn.config(state="normal")

if __name__ == "__main__":
    app = App()
    app.mainloop()